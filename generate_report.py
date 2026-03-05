"""
generate_report.py
Generates the complete Major Project Report as a .docx file
following the St Joseph's University template format.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re

# ── Paths ──────────────────────────────────────────────────────────────────
BASE_DIR      = r"C:\Users\dranu\OneDrive\Desktop\Osteoporosis Knee X ray"
BRAIN_DIR     = r"C:\Users\dranu\.gemini\antigravity\brain\43a56f5b-4689-42c1-aa22-606cea22d845"
OUTPUT_DOCX   = os.path.join(BASE_DIR, "Major_Project_Report_ODS.docx")
API_HTML      = os.path.join(BASE_DIR, "api_test_report.html")
CM_IMG        = os.path.join(BASE_DIR, "backend", "checkpoints", "confusion_matrices", "confusion_matrix_test.png")
ROC_IMG       = os.path.join(BASE_DIR, "backend", "checkpoints", "roc", "roc_all_classes.png")
ARCH_IMG      = os.path.join(BRAIN_DIR, "system_architecture_1772685633066.png")
PIPE_IMG      = os.path.join(BRAIN_DIR, "preprocessing_pipeline_1772685647660.png")
DFD0_IMG      = os.path.join(BRAIN_DIR, "dfd_level0_1772685661809.png")
DFD1_IMG      = os.path.join(BRAIN_DIR, "dfd_level1_1772685675282.png")

doc = Document()

# ── Styles ─────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

for s_name in ['Heading 1','Heading 2','Heading 3']:
    s = doc.styles[s_name]
    s.font.name = 'Times New Roman'
    s.font.bold = True
    s.font.color.rgb = RGBColor(0,0,0)
    if s_name == 'Heading 1': s.font.size = Pt(16)
    elif s_name == 'Heading 2': s.font.size = Pt(14)
    else: s.font.size = Pt(12)

def add_heading(text, level=1, center=False):
    h = doc.add_heading(text, level=level)
    if center:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h

def add_para(text='', bold=False, center=False, size=12, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_image(path, width=Inches(5.5), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(10)
    else:
        add_para(f"[Image not found: {os.path.basename(path)}]", italic=True)

def add_table_from_data(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return t

def page_break():
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
add_para('DEPARTMENT OF COMPUTER SCIENCE', bold=True, center=True, size=14)
doc.add_paragraph()
add_para('MAJOR PROJECT REPORT', bold=True, center=True, size=16)
doc.add_paragraph()
add_para('OSTEOPOROSIS DETECTION SYSTEM USING ENSEMBLE DEEP LEARNING', bold=True, center=True, size=14)
doc.add_paragraph()
doc.add_paragraph()
add_para('ST JOSEPH\'S UNIVERSITY', bold=True, center=True, size=14)
add_para('BANGALORE – 560027', bold=True, center=True, size=12)
add_para('www.sju.edu.in', center=True)
doc.add_paragraph()
add_para('2025–2026', bold=True, center=True, size=12)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CERTIFICATE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('CERTIFICATE', 1, center=True)
doc.add_paragraph()
cert_text = (
    "This is to certify that Smt. / Sri. ________________________________ has "
    "satisfactorily completed the Major Project titled \"Osteoporosis Detection "
    "System using Ensemble Deep Learning\" prescribed by St Joseph's University "
    "for the VI Semester degree course in BCA for the year 2025–26."
)
add_para(cert_text)
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_table(rows=2, cols=2)
t.cell(0,0).text = 'Signature of the Guide'
t.cell(0,1).text = 'Head of the Department'
t.cell(1,0).text = 'Examiner 1: ______________'
t.cell(1,1).text = 'Name: _____________________'
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('ACKNOWLEDGEMENT', 1, center=True)
doc.add_paragraph()
add_para(
    "I would like to express my sincere gratitude to the Department of Computer "
    "Science, St Joseph's University, Bangalore, for providing me with the "
    "opportunity to undertake this Major Project."
)
add_para(
    "I am deeply thankful to my project guide for their invaluable guidance, "
    "encouragement, and continuous support throughout the development of this project."
)
add_para(
    "I extend my gratitude to the Head of the Department and all faculty members "
    "for their support and for providing the necessary resources and infrastructure "
    "for the successful completion of this project."
)
add_para(
    "Finally, I am grateful to my family for their unwavering support and "
    "encouragement throughout my academic journey."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('TABLE OF CONTENTS', 1, center=True)
toc_data = [
    ('1', 'About the University'), ('2', 'Department of Computer Science'),
    ('3', 'Introduction'), ('4', 'Synopsis'), ('5', 'DFD / ER Diagrams'),
    ('6', 'Codes & Screenshots'), ('7', 'API Test Report'), ('8', 'Conclusion'),
]
t = doc.add_table(rows=1+len(toc_data), cols=3)
t.style = 'Table Grid'
for i, h in enumerate(['SL NO', 'Content', 'Page No']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for ri, (num, content) in enumerate(toc_data):
    t.rows[ri+1].cells[0].text = num
    t.rows[ri+1].cells[1].text = content
    t.rows[ri+1].cells[2].text = ''
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABOUT THE UNIVERSITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('ABOUT THE UNIVERSITY', 1, center=True)
add_para(
    "St Joseph's University (SJU) is a Jesuit university at the heart of Bengaluru, "
    "the Silicon City of India. Established in 1882 by Paris Foreign Fathers, the "
    "management of the college was handed over to the Jesuit order (Society of Jesus) "
    "in 1937. In February 2021, St Joseph's University bill was presented in the "
    "Karnataka Legislative Assembly. The college received its University status on "
    "2nd July 2022 and was inaugurated as India's first Public-Private Partnership "
    "University by the President of India, Smt. Droupadi Murmu on 27 September 2022."
)
add_heading('Milestones', 2)
milestones = [
    '1882 — Established by Paris Foreign Fathers',
    '1937 — Management handed over to the Jesuits (Society of Jesus)',
    '1986 — First affiliated college in Karnataka to offer postgraduate courses',
    '1998 — First college in Karnataka to get a research centre',
    '2005 — Granted autonomous status',
    '02 July 2022 — Received University status',
    '27 September 2022 — Inaugurated as India\'s first Public-Private Partnership University',
]
for m in milestones:
    p = doc.add_paragraph(m, style='List Bullet')
add_heading('Vision', 2)
add_para(
    "To form women and men for and with others, who through holistic education, "
    "strive for a just, secular, democratic, and ecologically sensitive society "
    "which empowers the poor, the oppressed and the marginalized."
)
add_heading('Mission', 2)
add_para(
    "In keeping with the Jesuit heritage, the university aims at an integral formation "
    "of the staff and the students, to be men and women who will be agents of societal "
    "change, enabling them to attain academic and human excellence."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEPARTMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('DEPARTMENT OF COMPUTER SCIENCE', 1, center=True)
add_para(
    "The Department was founded in 1986, offering BSc in Computer Science along with "
    "Physics and Mathematics. BCA is an undergraduate programme for candidates wishing "
    "to delve into the world of Computer Languages. MSc Computer Science is a "
    "postgraduate programme for candidates who wish to probe deeper into the realm of "
    "computer science."
)
add_para(
    "All final-year students develop and complete a project for an organisation outside "
    "the College, ensuring sufficient exposure and preparedness for a challenging career "
    "in the field of Information Technology."
)
add_heading('Vision', 2)
add_para(
    "To ensure that students have a deep and analytical understanding of the field and "
    "to enable them to use their immense potential to enhance the quality of human life "
    "by contributing to the field of Technology, with Ethical responsibilities."
)
add_heading('Mission', 2)
add_para(
    "The Department offers quality education in the Science of Computing by providing "
    "courses with Computer Applications and by focusing on Personal, Emotional, and "
    "Character development."
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('INTRODUCTION', 1, center=True)
add_heading('Background', 2)
add_para(
    "Osteoporosis is a systemic skeletal disorder characterised by reduced bone mass and "
    "deterioration of bone tissue, leading to increased bone fragility and fracture risk. "
    "It is often called the 'silent disease' because bone loss occurs without symptoms "
    "until a fracture happens. According to the World Health Organization (WHO), "
    "osteoporosis affects an estimated 200 million women worldwide and causes more than "
    "8.9 million fractures annually."
)
add_para(
    "The knee joint is particularly susceptible to osteoporotic changes, and X-ray imaging "
    "remains the most widely accessible and cost-effective screening tool. However, manual "
    "interpretation requires experienced radiologists and is subject to inter-observer "
    "variability. Artificial intelligence, specifically deep learning-based computer vision, "
    "offers the potential to provide fast, consistent, and accurate screening support."
)
add_heading('Motivation', 2)
add_para(
    "Early detection of osteoporosis can significantly reduce fracture risk through timely "
    "intervention, including calcium supplementation, vitamin D therapy, and lifestyle "
    "modifications. Access to specialised radiologists is limited in many healthcare settings, "
    "particularly in developing countries. An AI-assisted tool that can analyse knee X-rays "
    "and flag high-risk patients would be of significant clinical value."
)
add_heading('Problem Statement', 2)
add_para(
    "To develop a web-based AI application capable of classifying knee X-ray images into "
    "three categories — Normal, Osteopenia (early bone loss), and Osteoporosis (advanced "
    "bone loss) — and providing a clinical risk assessment by incorporating patient "
    "demographic information such as age, sex, and vitamin deficiency history."
)
add_heading('Objectives', 2)
objectives = [
    "Develop an ensemble deep learning model combining ResNet-50, DenseNet-121, and EfficientNet-B0",
    "Implement CLAHE preprocessing to enhance bone density visibility in X-ray images",
    "Integrate clinical covariate adjustment using patient demographics for improved accuracy",
    "Deploy the model as a full-stack web application with FastAPI backend and HTML/CSS/JS frontend",
    "Achieve test accuracy exceeding 90% on a balanced held-out test dataset",
]
for o in objectives:
    doc.add_paragraph(o, style='List Bullet')
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SYNOPSIS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('SYNOPSIS', 1, center=True)
add_heading('Abstract', 2)
add_para(
    "This project presents an AI-powered web application for automated osteoporosis risk "
    "assessment from knee X-ray images. The system employs an ensemble of three pre-trained "
    "Convolutional Neural Networks (CNNs) — ResNet-50, DenseNet-121, and EfficientNet-B0 — "
    "trained on a balanced dataset of 560 augmented knee X-ray images across three classes: "
    "Normal, Osteopenia, and Osteoporosis."
)
add_para(
    "Key innovations include: (1) CLAHE-based preprocessing to enhance bone density contrast, "
    "(2) a Bayesian prior clinical adjustment mechanism using patient age, sex, and vitamin "
    "deficiency history to modulate CNN output probabilities, (3) Test-Time Augmentation (TTA) "
    "using 5 transforms to improve prediction robustness, and (4) a weighted soft-voting "
    "ensemble achieving 95.96% test accuracy on a held-out test set of 99 images."
)

add_heading('Technology Stack', 2)
add_table_from_data(
    ['Component', 'Technology'],
    [
        ('Programming Language', 'Python 3.11'),
        ('Deep Learning Framework', 'PyTorch 2.x, TorchVision'),
        ('CNN Models', 'ResNet-50, DenseNet-121, EfficientNet-B0'),
        ('Image Processing', 'OpenCV (CLAHE), Pillow'),
        ('Backend API', 'FastAPI, Uvicorn'),
        ('Frontend', 'HTML5, CSS3, JavaScript (Vanilla)'),
        ('Training Platform', 'Google Colab (NVIDIA T4 GPU)'),
    ]
)
doc.add_paragraph()

add_heading('Dataset', 2)
add_table_from_data(
    ['Split', 'Normal', 'Osteopenia', 'Osteoporosis', 'Total'],
    [
        ('Training', '230', '165', '165', '560'),
        ('Testing',  '33',  '33',  '33',  '99'),
        ('Total',    '263', '198', '198', '659'),
    ]
)
doc.add_paragraph()

add_heading('Model Architecture', 2)
add_table_from_data(
    ['Model', 'Ensemble Weight', 'Parameters', 'Key Feature'],
    [
        ('ResNet-50',       '40%', '25.6M', 'Residual skip connections'),
        ('DenseNet-121',    '35%', '8.0M',  'Dense feature reuse'),
        ('EfficientNet-B0', '25%', '5.3M',  'Compound scaling'),
    ]
)
doc.add_paragraph()

add_heading('System Architecture', 2)
add_image(ARCH_IMG, caption='Figure 1: System Architecture Diagram')
doc.add_paragraph()

add_heading('Preprocessing Pipeline', 2)
add_image(PIPE_IMG, caption='Figure 2: Image Preprocessing Pipeline')
doc.add_paragraph()

add_heading('Results', 2)
add_table_from_data(
    ['Class', 'Precision', 'Recall', 'F1-Score', 'Support'],
    [
        ('Normal',       '0.94', '0.97', '0.96', '33'),
        ('Osteopenia',   '0.94', '1.00', '0.97', '33'),
        ('Osteoporosis', '1.00', '0.91', '0.95', '33'),
        ('Weighted Avg', '0.96', '0.96', '0.96', '99'),
    ]
)
doc.add_paragraph()
add_para('Overall Test Accuracy: 95.96%', bold=True, center=True)

add_heading('Confusion Matrix', 2)
add_image(CM_IMG, caption='Figure 3: Confusion Matrix — New Ensemble (ResNet-50 + DenseNet-121 + EfficientNet-B0)')
doc.add_paragraph()

add_heading('ROC Curves', 2)
add_image(ROC_IMG, caption='Figure 4: ROC Curves for All Classes')
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DFD / ER DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('DFD / ER DIAGRAMS', 1, center=True)

add_heading('DFD Level 0 — Context Diagram', 2)
add_para(
    "The Level 0 DFD (Context Diagram) shows the system as a single process, with "
    "external entities and data flows at the highest level of abstraction."
)
add_image(DFD0_IMG, caption='Figure 5: DFD Level 0 — Context Diagram')
doc.add_paragraph()

add_heading('DFD Level 1 — System Decomposition', 2)
add_para(
    "The Level 1 DFD decomposes the system into four major sub-processes: "
    "Image Preprocessing, Ensemble Prediction, Clinical Adjustment, and Risk Assessment."
)
add_image(DFD1_IMG, caption='Figure 6: DFD Level 1 — System Decomposition')
doc.add_paragraph()

add_table_from_data(
    ['Process', 'Description'],
    [
        ('1.0 Image Preprocessing', 'Applies CLAHE, resizes to 384x384, normalises the X-ray image'),
        ('2.0 Ensemble Prediction', 'Runs ResNet-50, DenseNet-121, EfficientNet-B0 with TTA; weighted soft voting'),
        ('3.0 Clinical Adjustment', 'Applies Bayesian prior using patient age, sex, and vitamin deficiency'),
        ('4.0 Risk Assessment',     'Determines urgency level: Low / Moderate / High / Critical'),
    ]
)
page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CODES & SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('CODES & SCREENSHOTS', 1, center=True)

add_heading('Key Source Files', 2)
add_table_from_data(
    ['File', 'Purpose'],
    [
        ('backend/app.py',     'FastAPI server — defines /predict endpoint and CORS settings'),
        ('backend/model.py',   'Core inference — ensemble, TTA, clinical adjustment'),
        ('backend/dataset.py', 'Dataset class, data loaders, transforms, model definitions'),
        ('backend/train.py',   'Training pipeline with early stopping and checkpointing'),
        ('frontend/index.html','Main web page — form, upload zone, results display'),
        ('frontend/script.js', 'API calls, response rendering, drag-and-drop logic'),
        ('frontend/styles.css','Dark-themed responsive CSS styling'),
    ]
)
doc.add_paragraph()

snippets = [
    ('Code Snippet 1: CLAHE Preprocessing (dataset.py)',
"""class CLAHETransform:
    def __init__(self, clip_limit=2.0, tile_grid_size=(8, 8)):
        self.clip_limit = clip_limit
        self.tile_grid_size = tile_grid_size

    def __call__(self, pil_img):
        img_array = np.array(pil_img.convert("L"))
        clahe = cv2.createCLAHE(
            clipLimit=self.clip_limit,
            tileGridSize=self.tile_grid_size
        )
        enhanced = clahe.apply(img_array)
        enhanced_rgb = cv2.cvtColor(enhanced, cv2.COLOR_GRAY2RGB)
        return Image.fromarray(enhanced_rgb)"""),

    ('Code Snippet 2: Weighted Ensemble with TTA (model.py)',
"""ENSEMBLE_WEIGHTS = [0.40, 0.35, 0.25]

def tta_predict(model, x):
    probs = torch.zeros((1, 3)).to(device)
    with torch.no_grad():
        for tf in tta_transforms:
            probs += F.softmax(model(tf(x)), dim=1)
    return probs / len(tta_transforms)

def ensemble_predict(models, x):
    return sum(w * tta_predict(m, x)
               for m, w in zip(models, ENSEMBLE_WEIGHTS))"""),

    ('Code Snippet 3: Clinical Covariate Adjustment (model.py)',
"""def clinical_adjust(probs_tensor, age, sex, vitamin_def):
    probs = probs_tensor.detach().cpu().numpy().flatten()
    risk_mult = np.ones(3)
    if age >= 70:
        risk_mult[0] *= 0.70; risk_mult[2] *= 1.50
    if sex.lower() == "female" and age >= 65:
        risk_mult[2] *= 1.40
    if vitamin_def:
        risk_mult[2] *= 1.35
    adjusted = probs * risk_mult
    adjusted = adjusted / adjusted.sum()
    return torch.tensor(adjusted, dtype=torch.float32).unsqueeze(0)"""),

    ('Code Snippet 4: FastAPI Prediction Endpoint (app.py)',
"""@app.post("/predict")
async def predict(
    age: int = Form(...), sex: str = Form(...),
    vitamin_deficiency: bool = Form(...),
    xray: UploadFile = File(...)
):
    img = preprocess_image(xray.file)
    base_probs = ensemble_predict(models, img)
    adjusted_probs = clinical_adjust(base_probs, age, sex, vitamin_deficiency)
    pred_idx = adjusted_probs.argmax(1).item()
    prediction = classes[pred_idx]
    urgency, message = assess_risk(prediction, age, sex, vitamin_deficiency)
    return {"prediction": prediction, "urgency": urgency, ...}"""),
]

for title, code in snippets:
    add_heading(title, 2)
    p = doc.add_paragraph()
    runner = p.add_run(code)
    runner.font.name = 'Courier New'
    runner.font.size = Pt(9)
    doc.add_paragraph()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# API TEST REPORT
# ══════════════════════════════════════════════════════════════════════════════
add_heading('API TEST REPORT', 1, center=True)
add_para(
    "The API test suite validates all backend endpoints of the Osteoporosis Detection "
    "System. Tests were run against the new ensemble (ResNet-50 + DenseNet-121 + "
    "EfficientNet-B0) using the api_test.py script."
)
doc.add_paragraph()

add_heading('Test Summary', 2)
add_table_from_data(
    ['Metric', 'Value'],
    [
        ('Total Tests',       '12'),
        ('Tests Passed',      '12'),
        ('Tests Failed',      '0'),
        ('Pass Rate',         '100%'),
        ('Backend URL',       'http://127.0.0.1:8000'),
        ('Models Under Test', 'ResNet-50 + DenseNet-121 + EfficientNet-B0'),
    ]
)
doc.add_paragraph()

add_heading('Test Case Details', 2)
test_cases = [
    ('1',  'GET /',                          'Health Check',                          'PASS', '200'),
    ('2',  'POST /predict',                  'Valid Normal X-ray (Male, 45)',          'PASS', '200'),
    ('3',  'POST /predict',                  'Valid Osteopenia X-ray (Female, 62)',    'PASS', '200'),
    ('4',  'POST /predict',                  'Valid Osteoporosis X-ray (Female, 72)', 'PASS', '200'),
    ('5',  'POST /predict',                  'Missing X-ray file (validation error)', 'PASS', '422'),
    ('6',  'POST /predict',                  'Missing age field',                     'PASS', '422'),
    ('7',  'POST /predict',                  'Missing sex field',                     'PASS', '422'),
    ('8',  'POST /predict',                  'Invalid file type (text/plain)',         'PASS', '400'),
    ('9',  'POST /predict',                  'Elderly female (80) with vitamin def',  'PASS', '200'),
    ('10', 'POST /predict',                  'Young male (25), no risk factors',      'PASS', '200'),
    ('11', 'GET /predict',                   'Wrong HTTP method',                     'PASS', '405'),
    ('12', 'GET /nonexistent',               '404 Not Found',                         'PASS', '404'),
]
add_table_from_data(
    ['No.', 'Endpoint', 'Test Case', 'Result', 'HTTP Status'],
    test_cases
)
doc.add_paragraph()

add_heading('API Endpoint Specification', 2)
add_heading('GET /', 3)
add_para('Returns the health status and model information.')
p = doc.add_paragraph()
p.add_run('Response:\n').bold = True
code_p = doc.add_paragraph()
run = code_p.add_run('{"status": "ok", "model": "ensemble (ResNet50 + DenseNet121 + EfficientNet-B0)"}')
run.font.name = 'Courier New'
run.font.size = Pt(9)
doc.add_paragraph()

add_heading('POST /predict', 3)
add_para('Classifies a knee X-ray image and returns a risk assessment.')
add_table_from_data(
    ['Parameter', 'Type', 'Required', 'Description'],
    [
        ('xray',               'File (JPG/PNG)', 'Yes', 'Knee X-ray image'),
        ('age',                'Integer',        'Yes', 'Patient age (1-120)'),
        ('sex',                'String',         'Yes', '"Male" or "Female"'),
        ('vitamin_deficiency', 'Boolean',        'Yes', 'History of vitamin deficiency'),
    ]
)
doc.add_paragraph()
add_para('Sample Response:', bold=True)
code_p = doc.add_paragraph()
run = code_p.add_run('''{
  "prediction": "osteopenia",
  "confidence": 0.6234,
  "class_probabilities": {
    "normal": 0.1823,
    "osteopenia": 0.6234,
    "osteoporosis": 0.1943
  },
  "urgency": "High",
  "message": "High risk: post-menopausal age group.",
  "patient": {"age": 65, "sex": "Female", "vitamin_deficiency": true}
}''')
run.font.name = 'Courier New'
run.font.size = Pt(9)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('CONCLUSION', 1, center=True)
add_heading('Summary', 2)
add_para(
    "This project successfully developed and deployed an AI-powered web application for "
    "osteoporosis detection from knee X-ray images. The system leverages an ensemble of "
    "three state-of-the-art CNNs — ResNet-50, DenseNet-121, and EfficientNet-B0 — trained "
    "on 560 augmented knee X-ray images and evaluated on a balanced held-out test set of "
    "99 images, achieving a final test accuracy of 95.96%."
)

add_heading('Key Achievements', 2)
achievements = [
    "95.96% overall test accuracy — up from a baseline of ~74% with the previous AlexNet-based ensemble",
    "Perfect precision (1.00) for Osteoporosis class — zero false positives for the most critical condition",
    "Novel Bayesian prior clinical covariate adjustment using patient demographics",
    "Complete full-stack deployment: FastAPI REST API + responsive dark-themed web frontend",
    "12/12 API tests passing — 100% endpoint reliability",
    "Test-Time Augmentation (5 transforms) for more stable predictions",
]
for a in achievements:
    doc.add_paragraph(a, style='List Bullet')

add_heading('Limitations', 2)
limitations = [
    "Trained on a relatively small dataset (659 total images) — larger datasets would improve generalization",
    "The system is a screening aid, not a replacement for clinical diagnosis by qualified radiologists",
    "CPU inference takes ~2-3 seconds per image; GPU deployment would enable real-time performance",
]
for l in limitations:
    doc.add_paragraph(l, style='List Bullet')

add_heading('Future Work', 2)
future = [
    "Expand dataset using public sources (MURA, OAI, RSNA)",
    "Add Grad-CAM heatmap visualisation for model explainability",
    "Incorporate multi-view X-ray analysis (lateral + AP views)",
    "Deploy on cloud server (Render, Railway, or AWS) for public access",
    "Develop a mobile application for rural/remote healthcare settings",
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

add_heading('References', 2)
references = [
    "He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image Recognition. CVPR 2016.",
    "Huang, G., Liu, Z., et al. (2017). Densely Connected Convolutional Networks. CVPR 2017.",
    "Tan, M., & Le, Q. (2019). EfficientNet: Rethinking Model Scaling for CNNs. ICML 2019.",
    "Zuiderveld, K. (1994). Contrast Limited Adaptive Histogram Equalization. Graphics Gems IV.",
    "World Health Organization (2023). Osteoporosis and Musculoskeletal Conditions. WHO Report.",
    "FastAPI Documentation (2024). https://fastapi.tiangolo.com",
]
for i, r in enumerate(references, 1):
    doc.add_paragraph(f"{i}. {r}")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_DOCX)
print(f"\nReport saved to: {OUTPUT_DOCX}")
